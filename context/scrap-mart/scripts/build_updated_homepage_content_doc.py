from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "exports" / "Scrap_Mart_Updated_Homepage_Content.docx"


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    set_para(p)
    if bold_label:
      r = p.add_run(bold_label)
      set_run(r, bold=True)
      r = p.add_run(text)
      set_run(r)
    else:
      r = p.add_run(text)
      set_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_para(p, before=20, after=6)
        size = 20
        style = "Heading 1"
    elif level == 2:
        set_para(p, before=18, after=6)
        size = 16
        style = "Heading 2"
    else:
        set_para(p, before=16, after=4)
        size = 14
        style = "Heading 3"
    p.style = style
    r = p.add_run(text)
    set_run(r, size=size)
    return p


def add_items(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=4)
        r = p.add_run(item)
        set_run(r)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    set_para(title, before=0, after=3)
    run = title.add_run("Scrap Mart Updated Homepage Content")
    set_run(run, size=26)

    add_para(doc, "Updated for the homepage preview with branch-specific price navigation, Auckland and Christchurch phone CTAs, quote enquiry form, location imagery, and revised hero headline.")

    add_heading(doc, "Header and Navigation")
    add_para(doc, "Logo: Scrap Mart")
    add_para(doc, "Navigation: Prices, What we buy, Services, Customers, Locations, About, Contact, Blog")
    add_para(doc, "Prices dropdown: Auckland prices, Christchurch prices")
    add_para(doc, "Header CTAs: Auckland: 09 941 7162, Christchurch: 02041234569, Get Quote")
    add_para(doc, "The location line below the logo has been removed.")

    add_heading(doc, "Hero")
    add_para(doc, "Eyebrow: Cash for scrap metal disposal")
    add_para(doc, "H1: Get Cash for Scrap Metal Disposal")
    add_para(doc, "Scrap Mart buys scrap metal from tradies, businesses and households, with clear weekly prices, certified weighing and fast same-day payment through our Auckland and Christchurch yards. Drop off at Onehunga or Hillsborough, or talk to us about commercial pickup for larger volumes.")
    add_para(doc, "Hero CTAs: View scrap metal prices, Get a scrap metal quote")
    add_items(doc, [
        "Weekly prices by grade: Useful rates before you visit or enquire.",
        "Same-day payment: Cash or bank transfer after inspection.",
        "Certified weighing: Clear grading before final payment.",
        "Two branch locations: Onehunga and Hillsborough, Christchurch.",
    ])

    add_heading(doc, "Quote Form")
    add_items(doc, [
        "Name",
        "Phone or email",
        "Preferred location: Auckland - Onehunga, Christchurch - Hillsborough",
        "Type of scrap: Household scrap, Commercial scrap, Tradie scrap, Batteries, Copper or cable, Other metal",
        "Scrap details",
        "Submit CTA: Send enquiry",
    ])

    add_heading(doc, "Pricing Section")
    add_para(doc, "Eyebrow: Price-led enquiries")
    add_para(doc, "Heading: Check current scrap metal prices")
    add_para(doc, "Prices change with metal markets, grade, quantity and cleanliness. Scrap Mart shows weekly rates so customers can check common scrap metal prices before visiting the yard or requesting commercial pickup.")
    add_para(doc, "Price preview materials: Copper, Copper cable, Batteries, Aluminium, Brass, Electric motors")
    add_para(doc, "Price note: Final payment is confirmed after your material is inspected, weighed and graded at the yard. Separating clean materials from mixed or dirty materials can help you get a stronger return.")
    add_para(doc, "Disclaimer: Prices are indicative and subject to change without notice. Prices may vary by location, quantity, cleanliness and market conditions.")
    add_para(doc, "Pricing CTAs: View Auckland prices, View Christchurch prices")
    add_para(doc, "Section CTA: Need a price before you visit? Send your metal details or call the nearest Scrap Mart branch.")

    add_heading(doc, "Common Price Questions")
    add_heading(doc, "How do I know if I am getting the right price?", 3)
    add_para(doc, "Your final payment depends on the metal grade, weight, cleanliness and inspection at the yard. Scrap Mart weighs and grades material before payment, so clean separated copper, cable, aluminium, brass and batteries can be priced more accurately than mixed or contaminated loads.")
    add_heading(doc, "What if I am not sure what metal or grade I have?", 3)
    add_para(doc, "Call your nearest branch or send details before visiting. The team can help identify common grades such as bright copper, PVC copper wire, aluminium extrusion, stainless steel, brass, electric motors, lead batteries and mixed steel.")

    add_heading(doc, "Scrap Metal We Buy")
    add_para(doc, "Scrap Mart buys common and high-value scrap metals from trades, workshops, factories, demolition sites, households and regular commercial suppliers.")
    add_items(doc, [
        "Copper: Bright copper, No.1 copper, No.2 copper, domestic copper, copper radiators and other copper grades.",
        "Copper Wire And Cable: PVC copper wire, electrical cable, automotive wiring looms and wire grades based on copper recovery.",
        "Aluminium: Aluminium extrusion, wheels, clean sheet, cast aluminium, domestic aluminium, cans, swarf and radiators.",
        "Batteries And Lead: Lead car and industrial batteries, gel and AGM batteries, soft lead and lead wheel weights.",
        "Brass: Mixed brass, irony brass and brass swarf from trade, plumbing, workshop and commercial scrap loads.",
        "Stainless Steel: Stainless steel 304, stainless steel swarf and other stainless material accepted after grading.",
        "Steel And Iron: HMS insize, HMS oversize, clean light gauge steel, dirty light gauge steel and mixed steel loads.",
        "Motors And Compressors: Electric motors, starter motors, alternators, fridge compressors and related recoverable components.",
        "Radiators, Cylinders And Zinc: Radiators, air conditioners, heat pumps, hot water cylinders, clean zinc sheet and zinc anodes.",
    ])

    add_heading(doc, "Customers We Serve")
    add_para(doc, "Heading: Scrap metal options for households, commercial sellers and tradies")
    add_para(doc, "Choose the customer path that matches your scrap load, then check prices, visit the yard, or ask about collection for suitable commercial volumes.")
    add_para(doc, "Household Scrap: For your household scrap metal, use our free scrap metal drop-off service at 11 Edinburgh St, Onehunga. We will dispose or recycle your domestic scrap metal at no cost for total convenience.")
    add_para(doc, "Commercial Scrap: Get rid of your commercial scrap metal and make money at the same time. You can drop off at our Onehunga yard, or our reliable commercial scrap metal collection service is available for Auckland businesses.")
    add_para(doc, "Tradie Scrap: From roofing iron to copper piping and air con units, if you have tradie scrap metal that needs to be disposed of or recycled, Scrap Mart offers same-day payments - drop off or ask about our collection service.")

    add_heading(doc, "Services")
    add_para(doc, "Heading: Scrap metal services for businesses and tradies")
    add_para(doc, "Scrap Mart helps trade and commercial customers clear valuable scrap metal, organise suitable collections, and keep regular material moving through the yard.")
    add_items(doc, [
        "Scrap Metal Buyers: Sell accepted scrap metal by grade at Scrap Mart, with certified weighing and same-day payment after inspection.",
        "Scrap Metal Collection: Collection enquiries for suitable scrap volumes, depending on material type, quantity, access and location.",
        "Commercial Scrap Buyers: Support for businesses, workshops and repeat suppliers with regular copper, aluminium, battery, motor or mixed metal loads.",
        "Battery Recycling: Scrap Mart buys scrap batteries at both Auckland and Christchurch, including lead car, industrial, gel and AGM batteries.",
    ])

    add_heading(doc, "How to Sell Scrap Metal")
    add_items(doc, [
        "Check prices or ask the team: View current scrap metal prices online, call your nearest branch, or send details if you are unsure what grade you have.",
        "Drop off or request pickup: Bring your scrap metal to Onehunga or Hillsborough. Commercial pickup depends on material type, volume and location.",
        "Weigh and grade the material: Your scrap is inspected, weighed and graded so the final price matches the material, cleanliness and quantity.",
        "Get paid: Once the material is confirmed, Scrap Mart pays the agreed amount on the same day by cash or bank transfer.",
    ])

    add_heading(doc, "Locations")
    add_heading(doc, "Visit Scrap Mart in Auckland or Christchurch", 2)
    add_para(doc, "Choose the branch closest to you, or contact the team if you have commercial scrap that may need pickup.")
    add_para(doc, "Scrap Mart Onehunga: 11 Edinburgh Street, Onehunga, Auckland. Monday to Friday, 8:00am to 4:30pm. Phone: 09 941 7162.")
    add_para(doc, "Supports trade, commercial and household scrap metal customers from Onehunga, Penrose, Mt Roskill, Mount Eden, Greenlane, Mangere, Manukau, East Tamaki, South Auckland, New Lynn, Henderson, West Auckland and nearby areas.")
    add_para(doc, "Scrap Mart Hillsborough: 48 Curries Road, Hillsborough, Christchurch 8022. Monday to Friday, 8:00am to 4:30pm. Phone: 02041234569.")
    add_para(doc, "Supports scrap metal customers across Hillsborough, Woolston, Sydenham, Hornby, Bromley, Ferrymead, Linwood, Riccarton and wider Christchurch.")

    add_heading(doc, "Trust and FAQ")
    add_para(doc, "Why sell your scrap metal to Scrap Mart? Scrap metal prices matter, but so does trust. Scrap Mart focuses on clear grading, certified weighing, fast payment and practical service for people who bring in scrap regularly.")
    add_items(doc, [
        "Established scrap metal experience",
        "Clear weekly pricing",
        "Certified weighing",
        "Same-day payment",
        "Trade-friendly service",
        "Responsible metal recycling",
    ])
    add_para(doc, "Household scrap metal drop-off: Scrap Mart mainly supports trade and commercial scrap metal customers, but household customers can also drop off accepted metal items at the Auckland and Christchurch yards.")
    add_para(doc, "Free scrap metal drop-off for household appliances. Some items cannot be accepted, including asbestos, oil, lithium batteries, alkaline batteries, bio waste, gas bottles and flammable liquids.")
    add_para(doc, "Final CTA: Ready to sell your scrap metal? Check current scrap metal prices, visit your nearest Scrap Mart yard, or send an enquiry if you have commercial scrap metal that may need pickup.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
